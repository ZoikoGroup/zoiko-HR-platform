"""
scripts/seed_real_policy_docs.py
------------------------------------
Ingests the real, HR-provided policy documents from docs/ into the knowledge
base, extracting full text (paragraphs + tables) via python-docx. Idempotent
by title — safe to re-run; already-published titles are skipped.

Run: `python scripts/seed_real_policy_docs.py <organization_id> <owner_employee_id>`
"""

import datetime
import os
import sys
import time

sys.path.insert(0, ".")

import docx
from sqlalchemy.exc import OperationalError

from app.database import SessionLocal, engine
from app.modules.assistant import knowledge_service
from app.modules.assistant.models import KnowledgeSource, KnowledgeStatus

_REPO_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

REAL_DOCUMENTS = [
    {
        "title": "Employee Handbook & Code of Conduct",
        "source_type": "handbook",
        "path": "docs/HR_Chatbot_Employee_Handbook_Code_of_Conduct_Engineering_Wireframe.docx",
    },
    {
        "title": "Attendance & Leave Policy",
        "source_type": "policy",
        "path": "docs/Zoiko_HR_Chatbot_Attendance_Leave_Policy_Engineering_Wireframe_v2.0.docx",
    },
    {
        "title": "Expense & Reimbursement Policy",
        "source_type": "policy",
        "path": "docs/Zoiko_HR_Chatbot_Expense_Reimbursement_Policy_Engineering_Wireframe_v2.0.docx",
    },
    {
        "title": "Remote Work & WFH Policy",
        "source_type": "policy",
        "path": "docs/Zoiko_HR_Chatbot_Remote_Work_WFH_Policy_Engineering_Wireframe_v2.0.docx",
    },
    {
        "title": "Benefits & Compensation Policy",
        "source_type": "policy",
        "path": "docs/Zoiko_HR_Chatbot_Benefits_Compensation_Policy_Engineering_Wireframe_v2.0.docx",
    },
    {
        "title": "Onboarding & Offboarding Procedures",
        "source_type": "sop",
        "path": "docs/Zoiko_HR_Chatbot_Onboarding_Offboarding_Procedures_Engineering_Wireframe_v2.0.docx",
    },
    {
        "title": "Compliance & Jurisdiction Policy Documents",
        "source_type": "compliance",
        "path": "docs/Zoiko_HR_Chatbot_Compliance_Jurisdiction_Policy_Documents_Engineering_Wireframe_v2.0.docx",
    },
]


def extract_docx_text(path: str) -> str:
    d = docx.Document(path)
    parts = []
    for p in d.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in d.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))
    return "\n\n".join(parts)


def main():
    if len(sys.argv) != 3:
        print("Usage: python scripts/seed_real_policy_docs.py <organization_id> <owner_employee_id>")
        sys.exit(1)
    organization_id = int(sys.argv[1])
    owner_employee_id = int(sys.argv[2])

    db = SessionLocal()
    try:
        # One-time placeholder retirement already ran and is done (the
        # placeholder titles collided with 3 of the real document titles
        # below, which caused this step to incorrectly re-retire already-
        # correct real documents on a second run — removed rather than
        # fixed, since it has no further reason to run again).
        already_published = {
            s.title for s in db.query(KnowledgeSource)
            .filter(KnowledgeSource.organization_id == organization_id, KnowledgeSource.status == KnowledgeStatus.PUBLISHED)
            .all()
        }

        for doc in REAL_DOCUMENTS:
            if doc["title"] in already_published:
                print(f"SKIP (already published): {doc['title']}")
                continue
            full_path = os.path.join(_REPO_ROOT, doc["path"])
            content = extract_docx_text(full_path)
            print(f"Extracted {doc['path']}: {len(content)} chars")

            # A long embedding pass (large document -> many chunks) can run
            # long enough for Neon's serverless Postgres to drop an idle
            # connection mid-transaction, unrelated to query correctness.
            # Retry once on a fresh session/connection rather than losing
            # the whole run over one transient network hiccup.
            for attempt in (1, 2):
                try:
                    source = knowledge_service.create_source(
                        db, organization_id, owner_employee_id,
                        title=doc["title"], source_type=doc["source_type"], authority_tier="A",
                        content_text=content,
                        jurisdiction_code=None, worker_type=None, audience_role=None,
                        effective_from=datetime.date.today(), effective_to=None,
                    )
                    knowledge_service.publish_source(db, organization_id, source.id, owner_employee_id)
                    print(f"CREATED + PUBLISHED: {doc['title']} (source id {source.id})")
                    break
                except OperationalError as e:
                    print(f"Attempt {attempt} failed for {doc['title']}: {e}")
                    db.rollback()
                    db.close()
                    engine.dispose()  # drop the dead connection from the pool
                    if attempt == 2:
                        raise
                    time.sleep(3)
                    db = SessionLocal()
        print("\nDone.")
    finally:
        db.close()


if __name__ == "__main__":
    main()
